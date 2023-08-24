# Please install pandas, python docx, pretty table and termcolor before running this code to avoid errors
# Please install pandas, python docx, pretty table and termcolor before running this code to avoid errors

import pandas as pd
from docx import Document
from docx.shared import Cm
from prettytable import PrettyTable, MSWORD_FRIENDLY
from termcolor import colored

# Read course data from CSV file
df = pd.read_csv('Courses_Data.csv')

# Sort course data by day and start time
# df = df.sort_values(by=['start_time', 'day'])

# Define initial venue capacities
venues = {'KDLT': 100, 'NFLT': 250, 'CBN': 1000, 'FLT': 1200}

# Define empty dictionary to keep track of venue booking status for each day
venue_bookings = {day: {'CBN': [], 'KDLT': [], 'NFLT': [], 'FLT': []} for day in df['day'].unique()}

# Display a Menu for the user to add a new venue or proceed without adding one
print(colored("Welcome to Course Management System!",
          'yellow', attrs=['bold']))

# Define the menu options as a pretty table
table = PrettyTable()
table.field_names = [colored("Option", 'cyan'),
                         colored("Description", 'cyan')]
table.add_row([colored("1", 'green'), "Proceed without adding a new venue"])
table.add_row([colored("2", 'green'), "Add a new venue and capacity"])
print(table)
choice = int(input("Enter your choice (1 or 2): "))

# If the user chose to add a new venue and capacity, allow them to do so
if choice == 2:
    new_venue = input('Enter new venue name: ')
    new_capacity = int(input('Enter new venue capacity: '))
    venues[new_venue] = new_capacity

    # Update the venue_bookings dictionary to include the new venue
    for day in venue_bookings:
        venue_bookings[day].update({new_venue: []})

# Allocate venues for courses based on capacity and availability
for index, row in df.iterrows():
    # Get course details
    course = row['course']
    lecturer = row['supervisors']
    start_time = row['start_time']
    end_time = row['end_time']
    day = row['day']
    capacity = row['capacity']

    # Find available venues for this course
    available_venues = []
    for venue, venue_capacity in venues.items():
        if capacity <= venue_capacity:
            bookings_for_venue = venue_bookings[day][venue]
            if not bookings_for_venue or bookings_for_venue[-1]['end_time'] <= start_time:
                available_venues.append(venue)

    # Allocate venue for course
    if available_venues:
        venue = available_venues[0]
        df.loc[index, 'venue'] = venue
        venue_bookings[day][venue].append({'start_time': start_time, 'end_time': end_time})
    else:
        df.loc[index, 'venue'] = 'None'

# Create a Word document
doc = Document()

# Add heading
doc.add_heading('HACKATHON PROJECT (GROUP 5 (Participants))', level=1)
doc.add_heading('1. ANIAH MOSES LIPEUNIM (230890), 2. OGHENEKOHWO OGHENEMARO OGHENEVWOKE (230907), 3. SANGOGADE AYOMIDE EPHRAIM (223322), 4. OLUWATOLA ENOCH ADEBAYO (230919), 5. OLOWE ANTHONY OLUBOBA (230916)', level=1)
doc.add_heading('EXAMINATION TIMETABLE', level=1)

# Create a table with headers
table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'
header_row = table.rows[0].cells
header_row[0].text = 'Course'
header_row[1].text = 'Supervisors'
header_row[2].text = 'Start Time'
header_row[3].text = 'End Time'
header_row[4].text = 'Day'
header_row[5].text = 'Venue'

# Add data rows to the table
for _, row in df.iterrows():
    course = row['course']
    lecturer = row['supervisors']
    start_time = row['start_time']
    end_time = row['end_time']
    day = row['day']
    venue = row['venue']
    table_row = table.add_row().cells
    table_row[0].text = str(course)
    table_row[1].text = str(lecturer)
    table_row[2].text = str(start_time)
    table_row[3].text = str(end_time)
    table_row[4].text = str(day)
    table_row[5].text = str(venue)

# Set column widths for the table
widths = [Cm(3), Cm(3), Cm(3), Cm(3), Cm(2), Cm(2)]
for i, width in enumerate(widths):
    table.columns[i].width = width

# Save the document
doc.save('exam_timetable.docx')

# Create a PrettyTable to display on the console
table = PrettyTable(['Course', 'Supervisors', 'Start Time', 'End Time', 'Day', 'Venue'])
table.set_style(MSWORD_FRIENDLY)

# Add data rows to the PrettyTable
for _, row in df.iterrows():
    course = row['course']
    lecturer = row['supervisors']
    start_time = row['start_time']
    end_time = row['end_time']
    day = row['day']
    venue = row['venue']
    table.add_row([course, lecturer, start_time, end_time, day, venue])

# Set table properties
table.align = 'c'
table.padding_width = 1
table.header_color = 'cyan'
table.border = True
table.title = 'EXAMINATION TIMETABLE'

# Print the PrettyTable to console
print(table)
